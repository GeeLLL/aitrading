from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "INVESTOR_UPDATE_2026-07-17_CN.md"
OUTPUT = ROOT / "Robinhood_AI_Investor_Update_2026-07-17_CN.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17324D"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "EAF2F8"
GREEN = "176B3A"
RED = "9B1C1C"
BODY_FONT = "Arial Unicode MS"


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    tail = paragraph.add_run(" 页")
    set_font(tail, 9, color=MUTED)


def add_inline(paragraph, text, *, size=10.5, color="202124"):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size - 0.5, bold=True, color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_font(run, size, color=color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 9, 4),
    ):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("ROBINHOOD AI OPTIONS EXPERIMENT  |  INVESTOR UPDATE")
    set_font(run, 8.5, bold=True, color=MUTED)
    add_page_field(section.footer.paragraphs[0])


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("项目进展更新")
    set_font(run, 25, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Robinhood AI 自主期权交易实验")
    set_font(run, 14, color=DARK_BLUE)

    for label, value in (
        ("日期", "2026 年 7 月 17 日"),
        ("当前阶段", "历史无未来信息测试完成，准备进入实时 Paper / Shadow Mode"),
        ("账户范围", "$300 独立实验账户；真实下单权限保持关闭"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}：")
        set_font(r, 10.5, bold=True, color=NAVY)
        r = p.add_run(value)
        set_font(r, 10.5, color="202124")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    pBdr.append(bottom)
    pPr.append(pBdr)

    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("投资人摘要")
    set_font(r, 11, bold=True, color=NAVY)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, "今天完成了官方数据链路、硬风控、实时只读验证及 29 个交易日历史实验。原始机械策略未通过；严格低频过滤值得继续验证；AI 的独立增量价值尚未得到证明。项目具备进入 Paper / Shadow Mode 的工程基础，但不进入真实交易。", size=10.5)


def add_markdown_body(doc, lines):
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("| "):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [x.strip() for x in lines[i].strip().strip("|").split("|")]
                if not all(set(x) <= {"-", ":"} for x in cells):
                    rows.append(cells)
                i += 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                widths = [2500] + [(9360 - 2500) // (len(rows[0]) - 1)] * (len(rows[0]) - 1)
                widths[-1] += 9360 - sum(widths)
                set_table_geometry(table, widths)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, value in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        if r_idx == 0:
                            shade(cell, LIGHT_GRAY)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_after = Pt(0)
                        add_inline(p, value, size=9.3, color=NAVY if r_idx == 0 else "202124")
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
            continue
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, line, size=11, color=NAVY)
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1


def build():
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)
    add_masthead(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    # Skip source title/date/stage because the masthead already carries them.
    start = next(i for i, line in enumerate(lines) if line.startswith("## 一、"))
    add_markdown_body(doc, lines[start:])

    core = doc.core_properties
    core.title = "Robinhood AI 自主期权交易实验：投资人项目进展更新"
    core.subject = "2026-07-17 项目进展、历史实验结果与下一步计划"
    core.author = "Robinhood AI Trading Experiment"
    core.keywords = "Robinhood, AI, options, shadow trading, investor update"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
